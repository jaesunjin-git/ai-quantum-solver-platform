# file_service.py — v2.0 (Final)
# ============================================================

import os
import re
import pathlib
import asyncio
import logging
from typing import Optional

import pandas as pd
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

# 업로드 기본 경로
BASE_UPLOAD_DIR = pathlib.Path("uploads").resolve()

# 파일별 요약 최대 길이
MAX_SUMMARY_LENGTH = 3000


# ============================================================
# 1. 보안: 안전한 업로드 경로 생성
# ============================================================
def _get_safe_upload_dir(project_id: str) -> pathlib.Path:
    """Path Traversal 방지를 위한 경로 검증"""
    safe_id = re.sub(r"[^a-zA-Z0-9_\-]", "", str(project_id))
    if not safe_id:
        raise ValueError(f"Invalid project_id: {project_id}")

    upload_dir = (BASE_UPLOAD_DIR / safe_id).resolve()

    if not str(upload_dir).startswith(str(BASE_UPLOAD_DIR)):
        raise ValueError(f"Path traversal detected: {project_id}")

    return upload_dir


# ============================================================
# 2. 인코딩 안전 텍스트 읽기
# ============================================================
def _read_text_safe(file_path: str, max_chars: int = 2000) -> str:
    """UTF-8 → CP949 순서로 시도하여 텍스트를 읽음"""
    for encoding in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read(max_chars)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "(텍스트를 읽을 수 없는 인코딩입니다)"


# ============================================================
# 3. 파일 타입별 분석 핸들러
# ============================================================
def _analyze_csv(file_path: str, filename: str) -> str:
    """CSV 파일 분석"""
    try:
        df = pd.read_csv(file_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="cp949")

    summary = (
        f"### 📄 CSV: {filename}\n"
        f"- **Rows**: {len(df):,}\n"
        f"- **Columns** ({len(df.columns)}): {list(df.columns)}\n"
    )

    # 데이터 타입 정보
    dtype_info = df.dtypes.value_counts().to_dict()
    dtype_str = ", ".join(
        [f"{str(k)}: {v}개" for k, v in dtype_info.items()]
    )
    summary += f"- **Data Types**: {dtype_str}\n"

    # 결측치 정보
    null_counts = df.isnull().sum()
    cols_with_nulls = null_counts[null_counts > 0]
    if len(cols_with_nulls) > 0:
        null_str = ", ".join(
            [f"{col}({cnt})" for col, cnt in cols_with_nulls.items()]
        )
        summary += f"- **Missing Values**: {null_str}\n"
    else:
        summary += "- **Missing Values**: 없음\n"

    # 수치형 컬럼 기본 통계
    numeric_cols = df.select_dtypes(include="number")
    if not numeric_cols.empty:
        stats = numeric_cols.describe().round(2)
        summary += f"\n**Numeric Statistics:**\n{stats.to_markdown()}\n"

    # 미리보기
    summary += f"\n**Preview (Top 3):**\n{df.head(3).to_markdown(index=False)}\n"

    return summary


def _analyze_excel(file_path: str, filename: str) -> str:
    """Excel 파일 분석 (.xlsx / .xls 분기)"""
    engine = "openpyxl" if filename.endswith(".xlsx") else "xlrd"

    try:
        xls = pd.ExcelFile(file_path, engine=engine)
    except Exception as e:
        # 엔진 호환 실패 시 대체 시도
        fallback_engine = "xlrd" if engine == "openpyxl" else "openpyxl"
        try:
            xls = pd.ExcelFile(file_path, engine=fallback_engine)
            engine = fallback_engine
        except Exception:
            raise e

    sheet_names = xls.sheet_names

    summary = (
        f"### 📄 Excel: {filename}\n"
        f"- **Sheets ({len(sheet_names)})**: {', '.join(sheet_names)}\n"
    )

    for sheet in sheet_names:
        df = pd.read_excel(
            xls, sheet_name=sheet, header=None, engine=engine
        )
        real_rows = df.dropna(how="all").shape[0]
        real_cols = df.dropna(axis=1, how="all").shape[1]

        summary += (
            f"\n**[Sheet: {sheet}]**\n"
            f"- Active Rows: {real_rows:,}, "
            f"Active Cols: {real_cols:,}\n"
        )

        # 미리보기 (상위 5행)
        if len(df) > 0:
            preview = df.head(5).fillna("").to_string(
                index=False, header=False
            )
            summary += f"- Content:\n```\n{preview}\n```\n"

    return summary


def _analyze_pdf(file_path: str, filename: str) -> str:
    """PDF 파일 분석"""
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)

    text_content = ""
    for page in reader.pages[:3]:
        extracted = page.extract_text()
        if extracted:
            text_content += extracted + "\n"

    if not text_content.strip():
        text_content = "(텍스트를 추출할 수 없습니다. 이미지 기반 PDF일 수 있습니다.)"

    return (
        f"### 📄 PDF: {filename}\n"
        f"- **Pages**: {total_pages}\n"
        f"- **Content (first 3 pages):**\n"
        f"{text_content[:800]}..."
    )


def _analyze_docx(file_path: str, filename: str) -> str:
    """Word 파일 분석"""
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    full_text = "\n".join(paragraphs)

    # 테이블 정보
    table_info = ""
    if doc.tables:
        table_info = f"- **Tables**: {len(doc.tables)}개\n"

    return (
        f"### 📄 Word: {filename}\n"
        f"- **Paragraphs**: {len(paragraphs)}개\n"
        f"{table_info}"
        f"- **Content:**\n{full_text[:800]}..."
    )


def _analyze_txt(file_path: str, filename: str) -> str:
    """텍스트 파일 분석"""
    content = _read_text_safe(file_path)
    return (
        f"### 📄 Text: {filename}\n"
        f"- **Content:**\n{content}..."
    )


# ============================================================
# 4. 파일 타입 → 핸들러 매핑
# ============================================================
FILE_HANDLERS = {
    ".csv":  _analyze_csv,
    ".xlsx": _analyze_excel,
    ".xls":  _analyze_excel,
    ".pdf":  _analyze_pdf,
    ".docx": _analyze_docx,
    ".txt":  _analyze_txt,
}

SUPPORTED_EXTENSIONS = set(FILE_HANDLERS.keys())

# ============================================================
# 5-A. 정확한 팩트 데이터 추출 (LLM 없이 코드로 계산)
# ============================================================
def extract_data_facts(project_id: str) -> dict:
    """
    업로드된 파일에서 정확한 팩트 데이터를 추출.
    이 데이터는 LLM을 거치지 않고 코드로 계산된 확정값.
    """
    upload_dir = _get_safe_upload_dir(project_id)
    if not upload_dir.exists():
        return {"files": [], "total_records": 0, "error": "업로드된 파일 없음"}

    facts = {
        "files": [],
        "total_records": 0,
        "total_columns": 0,
        "all_columns": {},       # {파일명: [컬럼목록]}
        "unique_counts": {},     # {파일명.컬럼명: 고유값수}
        "missing_values": {},    # {파일명.컬럼명: 결측수}
        "sheet_info": {},        # {파일명: {시트명: {rows, cols}}}
    }

    filenames = sorted(os.listdir(str(upload_dir)))

    for filename in filenames:
        file_path = str(upload_dir / filename)
        if filename.startswith("."):
            continue

        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        file_fact = {"name": filename, "type": ext, "records": 0, "columns": []}

        try:
            if ext == ".csv":
                try:
                    df = pd.read_csv(file_path, encoding="utf-8")
                except UnicodeDecodeError:
                    df = pd.read_csv(file_path, encoding="cp949")

                file_fact["records"] = len(df)
                file_fact["columns"] = list(df.columns)
                facts["total_records"] += len(df)
                facts["total_columns"] += len(df.columns)
                facts["all_columns"][filename] = list(df.columns)

                for col in df.columns:
                    key = f"{filename}.{col}"
                    facts["unique_counts"][key] = int(df[col].nunique())
                    null_count = int(df[col].isnull().sum())
                    if null_count > 0:
                        facts["missing_values"][key] = null_count

            elif ext in (".xlsx", ".xls"):
                engine = "openpyxl" if ext == ".xlsx" else "xlrd"
                try:
                    xls = pd.ExcelFile(file_path, engine=engine)
                except Exception:
                    xls = pd.ExcelFile(file_path, engine="xlrd" if engine == "openpyxl" else "openpyxl")

                sheet_info = {}
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet, engine=engine if hasattr(xls, '_engine') else None)
                    rows = len(df)
                    cols = list(df.columns)

                    sheet_info[sheet] = {"rows": rows, "cols": len(cols), "column_names": cols}
                    file_fact["records"] += rows
                    facts["total_records"] += rows
                    facts["all_columns"][f"{filename}:{sheet}"] = cols

                    for col in df.columns:
                        key = f"{filename}:{sheet}.{col}"
                        facts["unique_counts"][key] = int(df[col].nunique())
                        null_count = int(df[col].isnull().sum())
                        if null_count > 0:
                            facts["missing_values"][key] = null_count

                facts["sheet_info"][filename] = sheet_info
                file_fact["columns"] = list(sheet_info.keys())

            elif ext == ".txt":
                content = _read_text_safe(file_path)
                file_fact["records"] = content.count("\n") + 1
                file_fact["type"] = "text"

            elif ext == ".pdf":
                reader = PdfReader(file_path)
                file_fact["records"] = len(reader.pages)
                file_fact["type"] = "pdf"

        except Exception as e:
            logger.error(f"Fact extraction error [{filename}]: {e}")
            file_fact["error"] = str(e)

        facts["files"].append(file_fact)

    return facts


async def extract_data_facts_async(project_id: str) -> dict:
    """비동기 래퍼"""
    return await asyncio.to_thread(extract_data_facts, project_id)

# ============================================================
# 5. 메인 분석 함수
# ============================================================
def _analyze_csv_summary_sync(project_id: str) -> str:
    """
    프로젝트 업로드 디렉토리의 모든 파일을 분석하여
    통합 요약 리포트를 생성하는 동기 함수.
    """
    upload_dir = _get_safe_upload_dir(project_id)

    if not upload_dir.exists():
        return "업로드된 파일이 없습니다. 먼저 파일을 업로드해 주세요."

    summary_report = []
    processed_count = 0
    error_count = 0

    # 파일명 정렬하여 일관된 순서 보장
    filenames = sorted(os.listdir(str(upload_dir)))

    for filename in filenames:
        file_path = str(upload_dir / filename)

        # 숨김 파일 건너뛰기
        if filename.startswith("."):
            continue

        # 확장자 추출
        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        # 지원하지 않는 파일 형식 건너뛰기
        handler = FILE_HANDLERS.get(ext)
        if handler is None:
            continue

        try:
            file_summary = handler(file_path, filename)

            # 길이 제한
            if len(file_summary) > MAX_SUMMARY_LENGTH:
                file_summary = (
                    file_summary[:MAX_SUMMARY_LENGTH]
                    + "\n...(Truncated)..."
                )

            summary_report.append(file_summary)
            processed_count += 1

        except Exception as e:
            error_count += 1
            error_msg = str(e)

            # 사용자 친화적 에러 메시지
            if "openpyxl" in error_msg:
                error_msg = (
                    "서버에 openpyxl 라이브러리가 설치되지 않았습니다. "
                    "(pip install openpyxl)"
                )
            elif "xlrd" in error_msg:
                error_msg = (
                    "서버에 xlrd 라이브러리가 설치되지 않았습니다. "
                    "(pip install xlrd)"
                )

            logger.error(f"File analysis error [{filename}]: {e}")
            summary_report.append(
                f"### ❌ {filename} 읽기 실패\n- 원인: {error_msg}"
            )

    # 결과가 없는 경우
    if not summary_report:
        return (
            "업로드된 파일 중 분석 가능한 파일이 없습니다. "
            f"(지원 형식: {', '.join(sorted(SUPPORTED_EXTENSIONS))})"
        )

    # 전체 요약 헤더 추가
    header = (
        f"## 📊 파일 분석 요약\n"
        f"- 분석 완료: {processed_count}개 파일\n"
    )
    if error_count > 0:
        header += f"- 분석 실패: {error_count}개 파일\n"
    header += "---\n"

    return header + "\n\n".join(summary_report)


async def analyze_csv_summary(project_id: str) -> str:
    """
    비동기 래퍼.
    동기 I/O(파일 읽기, pandas 파싱)를 별도 스레드에서 실행하여
    이벤트 루프 블로킹을 방지.
    """
    return await asyncio.to_thread(
        _analyze_csv_summary_sync, project_id
    )