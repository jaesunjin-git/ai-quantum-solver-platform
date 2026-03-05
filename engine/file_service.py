# file_service.py — v3.0 (TASK 4: Enhanced File Support + Constraint Data Extraction)
# ============================================================
#
# v3.0 변경사항:
#   - 신규 형식: .doc, .hwp, .hwpx, .tsv, .json, .ods
#   - PDF: 전체 페이지 텍스트 추출 (3페이지 제한 해제)
#   - DOCX: 테이블 내용 추출 강화
#   - 신규 함수: extract_constraint_relevant_data()
#   - 신규 함수: extract_full_text()
# ============================================================

import os
import re
import json
import pathlib
import asyncio
import logging
from typing import Dict, List, Optional, Any

import pandas as pd

logger = logging.getLogger(__name__)

# 업로드 기본 경로
BASE_UPLOAD_DIR = pathlib.Path("uploads").resolve()

# 파일별 요약 최대 길이
MAX_SUMMARY_LENGTH = 3000

# 텍스트 추출 시 전체 파일 최대 문자
MAX_FULL_TEXT_LENGTH = 50000


# ============================================================
# 1. 보안: 안전한 업로드 경로 생성
# ============================================================
def _get_safe_upload_dir(project_id: str) -> pathlib.Path:
    safe_id = re.sub(r"[^a-zA-Z0-9_\-]", "", str(project_id))
    if not safe_id:
        raise ValueError(f"Invalid project_id: {project_id}")
    upload_dir = (BASE_UPLOAD_DIR / safe_id).resolve()
    if not str(upload_dir).startswith(str(BASE_UPLOAD_DIR)):
        raise ValueError(f"Path traversal detected: {project_id}")
    return upload_dir


# ============================================================
# 2. 인코딩 안전 텍스트 읽기
# ============================================================
def _read_text_safe(file_path: str, max_chars: int = 2000) -> str:
    for encoding in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read(max_chars)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "(텍스트를 읽을 수 없는 인코딩입니다)"


def _read_text_full(file_path: str) -> str:
    """전체 텍스트를 읽되 MAX_FULL_TEXT_LENGTH로 제한"""
    for encoding in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read(MAX_FULL_TEXT_LENGTH)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ""


# ============================================================
# 3. 파일 타입별 분석 핸들러
# ============================================================

# ── CSV ──
def _analyze_csv(file_path: str, filename: str) -> str:
    try:
        df = pd.read_csv(file_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="cp949")

    summary = (
        f"### 📄 CSV: {filename}\n"
        f"- **Rows**: {len(df):,}\n"
        f"- **Columns** ({len(df.columns)}): {list(df.columns)}\n"
    )

    dtype_info = df.dtypes.value_counts().to_dict()
    dtype_str = ", ".join([f"{str(k)}: {v}개" for k, v in dtype_info.items()])
    summary += f"- **Data Types**: {dtype_str}\n"

    null_counts = df.isnull().sum()
    cols_with_nulls = null_counts[null_counts > 0]
    if len(cols_with_nulls) > 0:
        null_str = ", ".join([f"{col}({cnt})" for col, cnt in cols_with_nulls.items()])
        summary += f"- **Missing Values**: {null_str}\n"
    else:
        summary += "- **Missing Values**: 없음\n"

    numeric_cols = df.select_dtypes(include="number")
    if not numeric_cols.empty:
        stats = numeric_cols.describe().round(2)
        summary += f"\n**Numeric Statistics:**\n{stats.to_markdown()}\n"

    summary += f"\n**Preview (Top 3):**\n{df.head(3).to_markdown(index=False)}\n"
    return summary


# ── TSV ──
def _analyze_tsv(file_path: str, filename: str) -> str:
    try:
        df = pd.read_csv(file_path, sep="\t", encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, sep="\t", encoding="cp949")

    summary = (
        f"### 📄 TSV: {filename}\n"
        f"- **Rows**: {len(df):,}\n"
        f"- **Columns** ({len(df.columns)}): {list(df.columns)}\n"
    )

    null_counts = df.isnull().sum()
    cols_with_nulls = null_counts[null_counts > 0]
    if len(cols_with_nulls) > 0:
        null_str = ", ".join([f"{col}({cnt})" for col, cnt in cols_with_nulls.items()])
        summary += f"- **Missing Values**: {null_str}\n"
    else:
        summary += "- **Missing Values**: 없음\n"

    summary += f"\n**Preview (Top 3):**\n{df.head(3).to_markdown(index=False)}\n"
    return summary


# ── Excel ──
def _analyze_excel(file_path: str, filename: str) -> str:
    engine = "openpyxl" if filename.endswith(".xlsx") else "xlrd"
    try:
        xls = pd.ExcelFile(file_path, engine=engine)
    except Exception as e:
        fallback_engine = "xlrd" if engine == "openpyxl" else "openpyxl"
        try:
            xls = pd.ExcelFile(file_path, engine=fallback_engine)
            engine = fallback_engine
        except Exception:
            raise e

    sheet_names = xls.sheet_names
    summary = (
        f"### 📄 Excel: {filename}\n"
        f"- **Sheets ({len(sheet_names)})**: {', '.join(sheet_names)}\n"
    )

    for sheet in sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet, header=None, engine=engine)
        real_rows = df.dropna(how="all").shape[0]
        real_cols = df.dropna(axis=1, how="all").shape[1]
        summary += (
            f"\n**[Sheet: {sheet}]**\n"
            f"- Active Rows: {real_rows:,}, Active Cols: {real_cols:,}\n"
        )
        if len(df) > 0:
            preview = df.head(5).fillna("").to_string(index=False, header=False)
            summary += f"- Content:\n```\n{preview}\n```\n"

    return summary


# ── ODS (LibreOffice) ──
def _analyze_ods(file_path: str, filename: str) -> str:
    try:
        xls = pd.ExcelFile(file_path, engine="odf")
        sheet_names = xls.sheet_names
        summary = (
            f"### 📄 ODS: {filename}\n"
            f"- **Sheets ({len(sheet_names)})**: {', '.join(sheet_names)}\n"
        )
        for sheet in sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet, engine="odf")
            summary += (
                f"\n**[Sheet: {sheet}]**\n"
                f"- Rows: {len(df):,}, Cols: {len(df.columns)}\n"
                f"- Columns: {list(df.columns)}\n"
            )
            if len(df) > 0:
                summary += f"\n**Preview:**\n{df.head(3).to_markdown(index=False)}\n"
        return summary
    except ImportError:
        return (
            f"### 📄 ODS: {filename}\n"
            f"- ⚠️ odfpy 라이브러리가 필요합니다 (pip install odfpy)\n"
        )
    except Exception as e:
        return f"### 📄 ODS: {filename}\n- ❌ 읽기 실패: {e}\n"


# ── JSON ──
def _analyze_json(file_path: str, filename: str) -> str:
    text = _read_text_full(file_path)
    if not text:
        return f"### 📄 JSON: {filename}\n- ❌ 읽기 실패\n"

    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        return f"### 📄 JSON: {filename}\n- ❌ JSON 파싱 실패: {e}\n"

    summary = f"### 📄 JSON: {filename}\n"

    if isinstance(data, list):
        summary += f"- **Type**: Array ({len(data)} items)\n"
        if data and isinstance(data[0], dict):
            keys = list(data[0].keys())
            summary += f"- **Item keys**: {keys}\n"
            try:
                df = pd.DataFrame(data[:100])
                summary += f"\n**Preview:**\n{df.head(3).to_markdown(index=False)}\n"
            except Exception:
                summary += f"- **First item**: {json.dumps(data[0], ensure_ascii=False)[:300]}\n"
    elif isinstance(data, dict):
        summary += f"- **Type**: Object ({len(data)} keys)\n"
        summary += f"- **Keys**: {list(data.keys())[:20]}\n"
        preview = json.dumps(data, ensure_ascii=False, indent=2)[:500]
        summary += f"- **Preview**:\n```json\n{preview}\n```\n"
    else:
        summary += f"- **Type**: {type(data).__name__}\n"

    return summary


# ── PDF (전체 페이지) ──
def _analyze_pdf(file_path: str, filename: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return f"### 📄 PDF: {filename}\n- ⚠️ pypdf 라이브러리가 필요합니다\n"

    reader = PdfReader(file_path)
    total_pages = len(reader.pages)

    # 요약용은 처음 5페이지
    text_content = ""
    for page in reader.pages[:5]:
        extracted = page.extract_text()
        if extracted:
            text_content += extracted + "\n"

    if not text_content.strip():
        text_content = "(텍스트를 추출할 수 없습니다. 이미지 기반 PDF일 수 있습니다.)"

    return (
        f"### 📄 PDF: {filename}\n"
        f"- **Pages**: {total_pages}\n"
        f"- **Content (first 5 pages):**\n"
        f"{text_content[:1500]}...\n"
    )


# ── DOCX (테이블 포함) ──
def _analyze_docx(file_path: str, filename: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return f"### 📄 Word: {filename}\n- ⚠️ python-docx 라이브러리가 필요합니다\n"

    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    full_text = "\n".join(paragraphs)

    summary = (
        f"### 📄 Word: {filename}\n"
        f"- **Paragraphs**: {len(paragraphs)}개\n"
    )

    # 테이블 정보 강화
    if doc.tables:
        summary += f"- **Tables**: {len(doc.tables)}개\n"
        for i, table in enumerate(doc.tables[:3]):
            rows = len(table.rows)
            cols = len(table.columns)
            summary += f"  - Table {i+1}: {rows}행 x {cols}열\n"
            # 첫 2행 미리보기
            preview_rows = []
            for row in table.rows[:2]:
                cells = [cell.text.strip()[:30] for cell in row.cells]
                preview_rows.append(" | ".join(cells))
            if preview_rows:
                summary += f"    ```\n    {'    '.join(r + chr(10) for r in preview_rows)}    ```\n"

    summary += f"- **Content:**\n{full_text[:1200]}...\n"
    return summary


# ── DOC (레거시 Word) ──
def _analyze_doc(file_path: str, filename: str) -> str:
    """
    .doc 파일 처리.
    시도 순서: antiword → textract → python-docx (변환 후)
    """
    text = ""

    # 방법 1: antiword (리눅스/Mac에서 가장 신뢰성 높음)
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 방법 2: textract
    if not text:
        try:
            import textract
            raw = textract.process(file_path)
            text = raw.decode("utf-8", errors="replace")
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"textract failed for {filename}: {e}")

    # 방법 3: olefile로 직접 텍스트 추출 시도
    if not text:
        try:
            import olefile
            ole = olefile.OleFileIO(file_path)
            if ole.exists("WordDocument"):
                stream = ole.openstream("WordDocument")
                raw_data = stream.read()
                # 간단한 텍스트 추출 (바이너리에서 가독 문자만)
                text = raw_data.decode("utf-8", errors="replace")
                # 제어 문자 제거
                text = re.sub(r"[^\x20-\x7E\uAC00-\uD7A3\u3131-\u3163\u0000-\u0009\u000A\u000D]", "", text)
                text = re.sub(r"\s{3,}", "\n", text)
            ole.close()
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"olefile failed for {filename}: {e}")

    if text:
        return (
            f"### 📄 DOC (Legacy): {filename}\n"
            f"- **Content length**: {len(text)} chars\n"
            f"- **Content:**\n{text[:1200]}...\n"
        )
    else:
        return (
            f"### 📄 DOC (Legacy): {filename}\n"
            f"- ⚠️ .doc 파일을 읽을 수 없습니다.\n"
            f"  .docx로 변환하여 다시 업로드하시거나,\n"
            f"  antiword 또는 textract를 설치해 주세요.\n"
        )


# ── HWP / HWPX (한글) ──
def _analyze_hwp(file_path: str, filename: str) -> str:
    """
    한글 파일(.hwp, .hwpx) 처리.
    시도 순서: pyhwp → hwp5txt → olefile(hwp5) → hwpx(zip)
    """
    ext = pathlib.Path(filename).suffix.lower()
    text = ""

    # .hwpx는 ZIP 기반 XML
    if ext == ".hwpx":
        text = _extract_hwpx_text(file_path)
        if text:
            return (
                f"### 📄 HWPX: {filename}\n"
                f"- **Content length**: {len(text)} chars\n"
                f"- **Content:**\n{text[:1500]}...\n"
            )

    # .hwp: 방법 1 – pyhwp
    if not text:
        try:
            import hwp5
            from hwp5.hwp5txt import extract_text
            text = extract_text(file_path)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"pyhwp failed for {filename}: {e}")

    # .hwp: 방법 2 – olefile로 직접 스트림 읽기
    if not text:
        try:
            import olefile
            ole = olefile.OleFileIO(file_path)
            # HWP5 형식: PrvText 스트림에 미리보기 텍스트
            if ole.exists("PrvText"):
                stream = ole.openstream("PrvText")
                raw = stream.read()
                text = raw.decode("utf-16-le", errors="replace")
                text = text.replace("\x00", "")
            # BodyText 섹션들에서 추출 시도
            if not text.strip():
                for entry in ole.listdir():
                    path = "/".join(entry)
                    if path.startswith("BodyText/Section"):
                        try:
                            stream = ole.openstream(entry)
                            raw = stream.read()
                            # HWP 바이너리에서 텍스트 추출 (간이)
                            decoded = raw.decode("utf-16-le", errors="replace")
                            decoded = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", decoded)
                            text += decoded + "\n"
                        except Exception:
                            continue
            ole.close()
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"olefile HWP failed for {filename}: {e}")

    # .hwp: 방법 3 – subprocess (hwp5txt CLI)
    if not text:
        try:
            import subprocess
            result = subprocess.run(
                ["hwp5txt", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

    if text and text.strip():
        # 정리
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return (
            f"### 📄 HWP: {filename}\n"
            f"- **Content length**: {len(text)} chars\n"
            f"- **Content:**\n{text[:1500]}...\n"
        )
    else:
        return (
            f"### 📄 HWP: {filename}\n"
            f"- ⚠️ HWP 파일 텍스트 추출에 실패했습니다.\n"
            f"  pyhwp 설치를 권장합니다: pip install pyhwp\n"
            f"  또는 PDF/DOCX로 변환하여 업로드해 주세요.\n"
        )


def _extract_hwpx_text(file_path: str) -> str:
    """HWPX (ZIP 기반 XML) 텍스트 추출"""
    import zipfile
    try:
        text_parts = []
        with zipfile.ZipFile(file_path, "r") as zf:
            for name in zf.namelist():
                # Contents/section*.xml 에 본문 텍스트
                if "section" in name.lower() and name.endswith(".xml"):
                    try:
                        import xml.etree.ElementTree as ET
                        xml_content = zf.read(name).decode("utf-8")
                        root = ET.fromstring(xml_content)
                        # 모든 텍스트 노드 추출
                        for elem in root.iter():
                            if elem.text and elem.text.strip():
                                text_parts.append(elem.text.strip())
                            if elem.tail and elem.tail.strip():
                                text_parts.append(elem.tail.strip())
                    except Exception:
                        continue
        return "\n".join(text_parts)
    except zipfile.BadZipFile:
        return ""
    except Exception as e:
        logger.warning(f"HWPX extraction failed: {e}")
        return ""


# ── TXT ──
def _analyze_txt(file_path: str, filename: str) -> str:
    content = _read_text_safe(file_path)
    return (
        f"### 📄 Text: {filename}\n"
        f"- **Content:**\n{content}...\n"
    )


# ============================================================
# 4. 파일 타입 → 핸들러 매핑
# ============================================================
FILE_HANDLERS = {
    ".csv":  _analyze_csv,
    ".tsv":  _analyze_tsv,
    ".xlsx": _analyze_excel,
    ".xls":  _analyze_excel,
    ".ods":  _analyze_ods,
    ".json": _analyze_json,
    ".pdf":  _analyze_pdf,
    ".docx": _analyze_docx,
    ".doc":  _analyze_doc,
    ".hwp":  _analyze_hwp,
    ".hwpx": _analyze_hwp,
    ".txt":  _analyze_txt,
    ".md":   _analyze_txt,
    ".text": _analyze_txt,
}

SUPPORTED_EXTENSIONS = set(FILE_HANDLERS.keys())


# ============================================================
# 5. 전체 텍스트 추출 (제약조건 추출용)
# ============================================================
def extract_full_text(file_path: str, filename: str) -> str:
    """
    파일에서 가능한 모든 텍스트를 추출한다.
    제약조건 값 추출(LLM)에 사용할 컨텍스트 생성용.
    """
    ext = pathlib.Path(filename).suffix.lower()
    text = ""

    try:
        if ext == ".csv":
            try:
                df = pd.read_csv(file_path, encoding="utf-8")
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, encoding="cp949")
            text = df.to_string(max_rows=50, max_cols=20)

        elif ext == ".tsv":
            try:
                df = pd.read_csv(file_path, sep="\t", encoding="utf-8")
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, sep="\t", encoding="cp949")
            text = df.to_string(max_rows=50, max_cols=20)

        elif ext in (".xlsx", ".xls"):
            engine = "openpyxl" if ext == ".xlsx" else "xlrd"
            try:
                xls = pd.ExcelFile(file_path, engine=engine)
            except Exception:
                xls = pd.ExcelFile(file_path, engine="xlrd" if engine == "openpyxl" else "openpyxl")
            parts = []
            for sheet in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet)
                parts.append(f"[Sheet: {sheet}]\n{df.to_string(max_rows=30)}")
            text = "\n\n".join(parts)

        elif ext == ".ods":
            try:
                xls = pd.ExcelFile(file_path, engine="odf")
                parts = []
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet, engine="odf")
                    parts.append(f"[Sheet: {sheet}]\n{df.to_string(max_rows=30)}")
                text = "\n\n".join(parts)
            except Exception:
                pass

        elif ext == ".json":
            raw = _read_text_full(file_path)
            try:
                data = json.loads(raw)
                text = json.dumps(data, ensure_ascii=False, indent=2)[:MAX_FULL_TEXT_LENGTH]
            except Exception:
                text = raw

        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                parts = []
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        parts.append(extracted)
                text = "\n".join(parts)
            except Exception:
                pass

        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                # 테이블 내용도 포함
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))
                text = "\n".join(parts)
            except Exception:
                pass

        elif ext == ".doc":
            # _analyze_doc과 동일한 추출 로직 재사용
            try:
                import subprocess
                result = subprocess.run(
                    ["antiword", file_path], capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0:
                    text = result.stdout
            except Exception:
                pass
            if not text:
                try:
                    import olefile
                    ole = olefile.OleFileIO(file_path)
                    if ole.exists("PrvText"):
                        raw = ole.openstream("PrvText").read()
                        text = raw.decode("utf-16-le", errors="replace").replace("\x00", "")
                    ole.close()
                except Exception:
                    pass

        elif ext in (".hwp", ".hwpx"):
            if ext == ".hwpx":
                text = _extract_hwpx_text(file_path)
            if not text:
                try:
                    import olefile
                    ole = olefile.OleFileIO(file_path)
                    if ole.exists("PrvText"):
                        raw = ole.openstream("PrvText").read()
                        text = raw.decode("utf-16-le", errors="replace").replace("\x00", "")
                    if not text.strip():
                        for entry in ole.listdir():
                            path = "/".join(entry)
                            if path.startswith("BodyText/Section"):
                                try:
                                    raw = ole.openstream(entry).read()
                                    decoded = raw.decode("utf-16-le", errors="replace")
                                    decoded = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", decoded)
                                    text += decoded + "\n"
                                except Exception:
                                    continue
                    ole.close()
                except Exception:
                    pass

        elif ext in (".txt", ".md", ".text"):
            text = _read_text_full(file_path)

    except Exception as e:
        logger.error(f"extract_full_text failed [{filename}]: {e}")

    return text[:MAX_FULL_TEXT_LENGTH] if text else ""


# ============================================================
# 6. 제약조건 관련 데이터 추출
# ============================================================
def extract_constraint_relevant_data(
    project_id: str,
    constraint_name: str,
    detection_hints: dict,
    required_data_types: list,
) -> str:
    """
    특정 제약조건에 관련된 데이터 스니펫을 추출한다.
    문제 정의(Phase B)에서 LLM에 전달할 컨텍스트 생성용.

    Args:
        project_id: 프로젝트 ID
        constraint_name: 제약조건 이름 (예: 'max_driving_time')
        detection_hints: 탐지 키워드 (예: {'ko': ['최대', '승무시간'], 'en': ['max', 'driving']})
        required_data_types: 필요한 데이터 유형 (예: ['work_regulations'])

    Returns:
        관련 데이터 텍스트 스니펫 (LLM 프롬프트에 삽입용)
    """
    upload_dir = _get_safe_upload_dir(project_id)
    if not upload_dir.exists():
        return ""

    # 키워드 수집
    all_keywords = []
    if isinstance(detection_hints, dict):
        for lang_keywords in detection_hints.values():
            if isinstance(lang_keywords, list):
                all_keywords.extend(kw.lower() for kw in lang_keywords)
    elif isinstance(detection_hints, list):
        all_keywords.extend(h.lower() for h in detection_hints)

    if not all_keywords:
        return ""

    snippets = []
    max_snippet_length = 500

    for fp in sorted(upload_dir.iterdir()):
        if fp.is_dir() or fp.name.startswith("."):
            continue

        ext = fp.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        try:
            full_text = extract_full_text(str(fp), fp.name)
            if not full_text:
                continue

            # 키워드 매칭으로 관련 섹션 추출
            lines = full_text.split("\n")
            for i, line in enumerate(lines):
                line_lower = line.lower()
                if any(kw in line_lower for kw in all_keywords):
                    # 매칭된 줄 + 전후 2줄 컨텍스트
                    start = max(0, i - 2)
                    end = min(len(lines), i + 3)
                    context = "\n".join(lines[start:end])
                    if len(context) > max_snippet_length:
                        context = context[:max_snippet_length] + "..."
                    snippets.append(f"[{fp.name}] {context}")

        except Exception as e:
            logger.warning(f"Constraint data extraction failed [{fp.name}]: {e}")

    return "\n---\n".join(snippets[:5]) if snippets else ""


async def extract_constraint_relevant_data_async(
    project_id: str,
    constraint_name: str,
    detection_hints: dict,
    required_data_types: list,
) -> str:
    return await asyncio.to_thread(
        extract_constraint_relevant_data,
        project_id, constraint_name, detection_hints, required_data_types
    )


# ============================================================
# 7. 정확한 팩트 데이터 추출 (LLM 없이 코드로 계산)
# ============================================================
def extract_data_facts(project_id: str) -> dict:
    upload_dir = _get_safe_upload_dir(project_id)
    if not upload_dir.exists():
        return {"files": [], "total_records": 0, "error": "업로드된 파일 없음"}

    facts = {
        "files": [],
        "total_records": 0,
        "total_columns": 0,
        "all_columns": {},
        "unique_counts": {},
        "missing_values": {},
        "sheet_info": {},
    }

    filenames = sorted(os.listdir(str(upload_dir)))

    for filename in filenames:
        file_path = str(upload_dir / filename)
        if filename.startswith("."):
            continue

        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        file_fact = {"name": filename, "type": ext, "records": 0, "columns": []}

        try:
            if ext == ".csv":
                try:
                    df = pd.read_csv(file_path, encoding="utf-8")
                except UnicodeDecodeError:
                    df = pd.read_csv(file_path, encoding="cp949")
                file_fact["records"] = len(df)
                file_fact["columns"] = list(df.columns)
                facts["total_records"] += len(df)
                facts["total_columns"] += len(df.columns)
                facts["all_columns"][filename] = list(df.columns)
                for col in df.columns:
                    key = f"{filename}.{col}"
                    facts["unique_counts"][key] = int(df[col].nunique())
                    null_count = int(df[col].isnull().sum())
                    if null_count > 0:
                        facts["missing_values"][key] = null_count

            elif ext == ".tsv":
                try:
                    df = pd.read_csv(file_path, sep="\t", encoding="utf-8")
                except UnicodeDecodeError:
                    df = pd.read_csv(file_path, sep="\t", encoding="cp949")
                file_fact["records"] = len(df)
                file_fact["columns"] = list(df.columns)
                facts["total_records"] += len(df)
                facts["total_columns"] += len(df.columns)
                facts["all_columns"][filename] = list(df.columns)
                for col in df.columns:
                    key = f"{filename}.{col}"
                    facts["unique_counts"][key] = int(df[col].nunique())
                    null_count = int(df[col].isnull().sum())
                    if null_count > 0:
                        facts["missing_values"][key] = null_count

            elif ext in (".xlsx", ".xls"):
                engine = "openpyxl" if ext == ".xlsx" else "xlrd"
                try:
                    xls = pd.ExcelFile(file_path, engine=engine)
                except Exception:
                    xls = pd.ExcelFile(file_path, engine="xlrd" if engine == "openpyxl" else "openpyxl")
                sheet_info = {}
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet)
                    rows = len(df)
                    cols = list(df.columns)
                    sheet_info[sheet] = {"rows": rows, "cols": len(cols), "column_names": cols}
                    file_fact["records"] += rows
                    facts["total_records"] += rows
                    facts["all_columns"][f"{filename}:{sheet}"] = cols
                    for col in df.columns:
                        key = f"{filename}:{sheet}.{col}"
                        facts["unique_counts"][key] = int(df[col].nunique())
                        null_count = int(df[col].isnull().sum())
                        if null_count > 0:
                            facts["missing_values"][key] = null_count
                facts["sheet_info"][filename] = sheet_info
                file_fact["columns"] = list(sheet_info.keys())

            elif ext == ".ods":
                try:
                    xls = pd.ExcelFile(file_path, engine="odf")
                    sheet_info = {}
                    for sheet in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=sheet, engine="odf")
                        rows = len(df)
                        cols = list(df.columns)
                        sheet_info[sheet] = {"rows": rows, "cols": len(cols), "column_names": cols}
                        file_fact["records"] += rows
                        facts["total_records"] += rows
                        facts["all_columns"][f"{filename}:{sheet}"] = cols
                    facts["sheet_info"][filename] = sheet_info
                    file_fact["columns"] = list(sheet_info.keys())
                except Exception as e:
                    file_fact["error"] = f"ODS read error: {e}"

            elif ext == ".json":
                raw = _read_text_full(file_path)
                try:
                    data = json.loads(raw)
                    if isinstance(data, list):
                        file_fact["records"] = len(data)
                        if data and isinstance(data[0], dict):
                            file_fact["columns"] = list(data[0].keys())
                            facts["all_columns"][filename] = file_fact["columns"]
                            facts["total_records"] += len(data)
                    elif isinstance(data, dict):
                        file_fact["records"] = len(data)
                        file_fact["columns"] = list(data.keys())
                except Exception:
                    file_fact["type"] = "json_error"

            elif ext == ".txt":
                content = _read_text_safe(file_path)
                file_fact["records"] = content.count("\n") + 1
                file_fact["type"] = "text"

            elif ext == ".pdf":
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    file_fact["records"] = len(reader.pages)
                    file_fact["type"] = "pdf"
                except Exception:
                    file_fact["type"] = "pdf"
                    file_fact["records"] = 0

            elif ext in (".docx", ".doc", ".hwp", ".hwpx"):
                file_fact["type"] = "document"
                file_fact["records"] = 1

        except Exception as e:
            logger.error(f"Fact extraction error [{filename}]: {e}")
            file_fact["error"] = str(e)

        facts["files"].append(file_fact)

    return facts


async def extract_data_facts_async(project_id: str) -> dict:
    return await asyncio.to_thread(extract_data_facts, project_id)


# ============================================================
# 8. 메인 분석 함수
# ============================================================
def _analyze_csv_summary_sync(project_id: str) -> str:
    upload_dir = _get_safe_upload_dir(project_id)
    if not upload_dir.exists():
        return "업로드된 파일이 없습니다. 먼저 파일을 업로드해 주세요."

    summary_report = []
    processed_count = 0
    error_count = 0

    filenames = sorted(os.listdir(str(upload_dir)))

    for filename in filenames:
        file_path = str(upload_dir / filename)
        if filename.startswith("."):
            continue

        # 디렉토리 스킵
        if os.path.isdir(file_path):
            continue

        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        handler = FILE_HANDLERS.get(ext)
        if handler is None:
            continue

        try:
            file_summary = handler(file_path, filename)
            if len(file_summary) > MAX_SUMMARY_LENGTH:
                file_summary = file_summary[:MAX_SUMMARY_LENGTH] + "\n...(Truncated)..."
            summary_report.append(file_summary)
            processed_count += 1
        except Exception as e:
            error_count += 1
            error_msg = str(e)
            if "openpyxl" in error_msg:
                error_msg = "서버에 openpyxl 라이브러리가 필요합니다 (pip install openpyxl)"
            elif "xlrd" in error_msg:
                error_msg = "서버에 xlrd 라이브러리가 필요합니다 (pip install xlrd)"
            logger.error(f"File analysis error [{filename}]: {e}")
            summary_report.append(f"### ❌ {filename} 읽기 실패\n- 원인: {error_msg}")

    if not summary_report:
        return (
            "업로드된 파일 중 분석 가능한 파일이 없습니다. "
            f"(지원 형식: {', '.join(sorted(SUPPORTED_EXTENSIONS))})"
        )

    header = (
        f"## 📊 파일 분석 요약\n"
        f"- 분석 완료: {processed_count}개 파일\n"
    )
    if error_count > 0:
        header += f"- 분석 실패: {error_count}개 파일\n"
    header += "---\n"

    return header + "\n\n".join(summary_report)


async def analyze_csv_summary(project_id: str) -> str:
    return await asyncio.to_thread(_analyze_csv_summary_sync, project_id)
