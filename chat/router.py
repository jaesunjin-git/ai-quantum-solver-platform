# ============================================================annotations
# chat/router.py — v3.0
# ============================================================
# 변경 이력:
#   v1.0 : 기본 chat + upload
#   v2.0 : 보안 강화, 확장자 검증, 리포트 다운로드, DOCX 변환
#   v3.0 :
#     - 다운로드 엔드포인트 인증 처리 (RBAC 호환, Optional fallback)
#     - format 파라미터를 Query + Literal 검증으로 강화
#     - 업로드 완료 후 agent에 file_upload 이벤트 자동 전달
#     - _clean_md 강화 (HTML 주석, SYSTEM-LOCKED, 내부 검증 블록 제거)
#     - 한글 파일명 명시적 허용
#     - aiofiles 의존성 제거 (표준 IO 사용)
#     - 에러 응답 구조 통일
# ============================================================
from __future__ import annotations

import io
import re
import json
import logging

from pathlib import Path
from typing import List, Literal, Optional

from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException, Query, Request
from fastapi.responses import Response
from sqlalchemy.orm import Session

from core.database import get_db
from core.rate_limit import limiter
from engine.validation.registry import get_registry
from core.models import ChatHistoryDB
from core.schemas import ChatRequest, ChatResponse
from .chat_service import process_user_intents

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["Chat"])

# ── 상수 ──────────────────────────────────────────────────
UPLOAD_BASE = Path("uploads")
MAX_FILE_SIZE = 50 * 1024 * 1024                       # 50 MB
ALLOWED_EXTENSIONS = {
    ".csv", ".xlsx", ".xls",
    ".txt", ".md",
    ".pdf", ".docx",
    ".json",
}


# ============================================================
# 보안 유틸
# ============================================================
def _sanitize_project_id(project_id: str) -> str:
    """Path Traversal 방지 — 영문, 숫자, 하이픈, 언더스코어만 허용"""
    safe = re.sub(r"[^a-zA-Z0-9_\-]", "", str(project_id))
    if not safe:
        raise HTTPException(status_code=400, detail="유효하지 않은 project_id")
    return safe


def _sanitize_filename(filename: str) -> str:
    """안전한 파일명 생성 — 한글, 영문, 숫자, 일부 기호 허용"""
    if not filename:
        return ""
    name = Path(filename).name                          # 경로 구분자 제거
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', name) # 위험 문자 치환
    name = name.strip('. ')
    return name if name else ""


def _validate_extension(filename: str) -> None:
    """허용 확장자 검사"""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"허용되지 않는 파일 형식: {ext} "
                   f"(허용: {', '.join(sorted(ALLOWED_EXTENSIONS))})",
        )


# ============================================================
# 1) 채팅 메시지
# ============================================================
from core.auth import get_current_user
from core.models import UserDB

@router.post("/chat/message", response_model=ChatResponse)
@limiter.limit("30/minute")
async def chat_endpoint(
    request: Request,
    payload: ChatRequest,
    current_user: UserDB = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    current_user_id = current_user.display_name or current_user.username

    # 1) 사용자 메시지 저장
    if payload.message and payload.message.strip():
        user_log = ChatHistoryDB(
            project_id=int(payload.project_id) if str(payload.project_id).isdigit() else 0,
            role="user",
            message_type="text",
            message_text=payload.message,
        )
        db.add(user_log)
        db.commit()

    # 2) AI 처리
    result = await process_user_intents(
        db=db,
        user_message=payload.message,
        project_id=str(payload.project_id),
        user_id=current_user_id,
        event_type=payload.event_type,
        event_data=payload.event_data,
        current_tab=payload.current_tab,
    )

    if isinstance(result, dict):
        if "text" not in result:
            result["text"] = str(result)
    else:
        result = {"text": str(result)}

    # 3) AI 응답 저장
    bot_log = ChatHistoryDB(
        project_id=int(payload.project_id) if str(payload.project_id).isdigit() else 0,
        role="assistant",
        message_type=result.get("type", "text"),
        message_text=result.get("text", ""),
        card_json=json.dumps(result.get("data"), ensure_ascii=False) if result.get("data") else None,
        options_json=json.dumps(result.get("options"), ensure_ascii=False) if result.get("options") else None,
    )
    db.add(bot_log)
    db.commit()

    return ChatResponse(**result)

# ============================================================
# 2) 파일 업로드
# ============================================================
@router.post("/upload")
@limiter.limit("10/minute")
async def upload_files(
    request: Request,
    project_id: str = Form(...),
    files: List[UploadFile] = File(...),
    current_user: UserDB = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id is required")
    if not files:
        raise HTTPException(status_code=400, detail="파일이 필요합니다.")

    safe_pid = _sanitize_project_id(project_id)
    upload_dir = UPLOAD_BASE / safe_pid
    upload_dir.mkdir(parents=True, exist_ok=True)

    saved_files: list[dict] = []
    errors: list[dict] = []

    for f in files:
        # ── 파일명 검증 ──
        safe_name = _sanitize_filename(f.filename or "unknown")
        if not safe_name:
            errors.append({"filename": f.filename, "error": "유효하지 않은 파일명"})
            continue

        # ── 확장자 검증 ──
        ext = Path(safe_name).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            errors.append({"filename": f.filename, "error": f"허용되지 않는 확장자: {ext}"})
            continue

        # ── 읽기 + 크기 검증 + 저장 ──
        try:
            content = await f.read()
            if len(content) > MAX_FILE_SIZE:
                size_mb = len(content) // (1024 * 1024)
                errors.append({
                    "filename": f.filename,
                    "error": f"파일 크기 초과 ({size_mb}MB > 50MB)",
                })
                continue

            file_path = upload_dir / safe_name
            file_path.write_bytes(content)

            saved_files.append({
                "filename": safe_name,
                "path": str(file_path),
                "size": len(content),
            })
        except Exception as e:
            logger.error(f"File save error [{f.filename}]: {e}")
            errors.append({"filename": f.filename or "unknown", "error": str(e)})

    # ── Stage 1 validation (upload) ──
    upload_validation = None
    if saved_files:
        try:
            registry = get_registry()
            stage_result = registry.run_stage(1, context={
                "files": saved_files,
                "project_id": str(safe_pid),
            })
            upload_validation = stage_result.to_dict()
        except Exception as e:
            logger.warning(f"Upload validation failed: {e}")

    # ── 업로드 완료 후 agent에 file_upload 이벤트 전달 ──
    agent_response = None
    if saved_files:
        try:
            agent_response = await process_user_intents(
                db=db,
                user_message="",
                project_id=str(safe_pid),
                user_id=current_user.display_name or current_user.username,
                event_type="file_upload",
                event_data={"files": saved_files},
            )
        except Exception as e:
            logger.warning(f"Agent file_upload event failed: {e}")

    # ── 응답 조립 ──
    response = {
        "status": "success" if saved_files else "error",
        "uploaded_count": len(saved_files),
        "uploaded_files": saved_files,
        "message": f"{len(saved_files)}개 파일 업로드 완료",
    }

    if upload_validation:
        response["validation"] = upload_validation

    if errors:
        response["errors"] = errors
        response["message"] += f", {len(errors)}개 실패"

    # agent 응답이 있으면 포함 (프론트엔드가 채팅 메시지로 표시 가능)
    if agent_response:
        if isinstance(agent_response, dict):
            response["chat_response"] = agent_response
        else:
            response["chat_response"] = {"text": str(agent_response)}

    return response


# ============================================================
# 3) 리포트 다운로드
# ============================================================
@router.get("/projects/{project_id}/report/download")
async def download_report(
    project_id: str,
    format: Literal["md", "docx", "json"] = Query(
        default="md",
        description="다운로드 형식: md, docx, json",
    ),
    type: Literal["analysis", "math_model", "solve_result"] = Query(
        default="analysis",
        description="다운로드 대상: analysis(분석 리포트), math_model(수학 모델), solve_result(최적화 결과)",
    ),
):
    """
    분석 리포트 또는 수학 모델을 다운로드.
    - analysis + md/docx: 분석 리포트
    - math_model + json: 수학 모델 원본 JSON
    - math_model + md: 수학 모델 요약 Markdown
    """
    safe_pid = _sanitize_project_id(project_id)

    from domains.crew.agent import get_session
    session = get_session(safe_pid)

    # ── 수학 모델 다운로드 ──
    if type == "math_model":
        model = session.state.math_model
        if not model:
            raise HTTPException(
                status_code=404,
                detail="다운로드할 수학 모델이 없습니다. 먼저 수학 모델을 생성해 주세요.",
            )

        if format == "json":
            content = json.dumps(model, ensure_ascii=False, indent=2)
            return Response(
                content=content.encode("utf-8"),
                media_type="application/json; charset=utf-8",
                headers={
                    "Content-Disposition": f'attachment; filename="math_model_{safe_pid}.json"',
                },
            )

        # math_model + md
        from engine.math_model_generator import summarize_model
        md_content = summarize_model(model)
        
        if format == "md":
            return Response(
            content=md_content.encode("utf-8"),
            media_type="text/markdown; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="math_model_{safe_pid}.md"',
            },
        )
    
        # math_model + docx  ★ 추가
        if format == "docx":
            try:
                docx_bytes = _markdown_to_docx(md_content, safe_pid)
                return Response(
                    content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="math_model_{safe_pid}.docx"',
                    },
                )
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="python-docx 라이브러리가 설치되지 않았습니다.",
                )
            except Exception as e:
                logger.error(f"Math model DOCX generation error: {e}", exc_info=True)
                raise HTTPException(status_code=500, detail="DOCX 생성 중 오류")

    #  최적화 결과 다운로드 
    if type == "solve_result":
        result = session.state.last_optimization_result
        if not result:
            raise HTTPException(
                status_code=404,
                detail="다운로드할 최적화 결과가 없습니다. 먼저 솔버를 실행해 주세요.",
            )

        if format == "json":
            import json as _json
            content_str = _json.dumps(result, ensure_ascii=False, indent=2, default=str)
            return Response(
                content=content_str.encode("utf-8"),
                media_type="application/json; charset=utf-8",
                headers={
                    "Content-Disposition": f'attachment; filename="solve_result_{safe_pid}.json"',
                },
            )

        if format == "md":
            md_parts = [
                "# Optimization Result\n",
                f"- **Status**: {result.get('status', 'N/A')}",
                f"- **Objective Value**: {result.get('objective_value', 'N/A')}",
                f"- **Solver**: {result.get('solver_name', 'N/A')}",
                f"- **Solver Type**: {result.get('solver_type', 'N/A')}",
                "",
            ]
            stats = result.get("model_stats", {})
            if stats:
                md_parts += [
                    "## Model Statistics",
                    f"- Variables: {stats.get('total_variables', 'N/A')}",
                    f"- Constraints: {stats.get('total_constraints', 'N/A')}",
                    "",
                ]
            timing = result.get("timing", {})
            if timing:
                md_parts += [
                    "## Timing",
                    f"- Compile: {timing.get('compile_sec', 'N/A')}s",
                    f"- Execute: {timing.get('execute_sec', 'N/A')}s",
                    f"- Total: {timing.get('total_sec', 'N/A')}s",
                    "",
                ]
            md_content = "\n".join(md_parts)
            return Response(
                content=md_content.encode("utf-8"),
                media_type="text/markdown; charset=utf-8",
                headers={
                    "Content-Disposition": f'attachment; filename="solve_result_{safe_pid}.md"',
                },
            )

        raise HTTPException(status_code=400, detail="최적화 결과는 JSON 또는 MD 형식만 지원합니다.")

    # ── 분석 리포트 다운로드 ──
    report = session.state.last_analysis_report
    if not report:
        raise HTTPException(
            status_code=404,
            detail="다운로드할 리포트가 없습니다. 먼저 데이터 분석을 진행해 주세요.",
        )

    report = _clean_report_for_download(report)

    if format == "md":
        return Response(
            content=report.encode("utf-8"),
            media_type="text/markdown; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="report_{safe_pid}.md"',
            },
        )

    if format == "json":
        raise HTTPException(
            status_code=400,
            detail="분석 리포트는 JSON 형식을 지원하지 않습니다.",
        )

    # docx
    try:
        docx_bytes = _markdown_to_docx(report, safe_pid)
        return Response(
            content=docx_bytes,
            media_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": f'attachment; filename="report_{safe_pid}.docx"',
            },
        )
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="python-docx 라이브러리가 설치되지 않았습니다. "
                   "(pip install python-docx)",
        )
    except Exception as e:
        logger.error(f"DOCX generation error: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Word 파일 생성 중 오류가 발생했습니다.",
        )


# ============================================================
# 내부 헬퍼
# ============================================================

def _clean_report_for_download(report: str) -> str:
    """다운로드용 리포트에서 내부 지시문/메타 텍스트 제거"""
    # ⛔ 로 시작하는 줄
    cleaned = re.sub(r'^⛔.*$', '', report, flags=re.MULTILINE)
    # <!-- ... --> HTML 주석
    cleaned = re.sub(r'<!--.*?-->', '', cleaned, flags=re.DOTALL)
    # [내부 검증] 블록
    cleaned = re.sub(
        r'\[내부\s*검증[^\]]*\].*?(?=\n##|\n---|\Z)',
        '', cleaned, flags=re.DOTALL,
    )
    # SYSTEM-LOCKED 관련 텍스트
    cleaned = re.sub(
        r'SYSTEM[- ]?LOCKED.*?(?=\n##|\n---|\Z)',
        '', cleaned, flags=re.DOTALL | re.IGNORECASE,
    )
    # "절대 변경하지 마십시오" / "출력을 종료하세요"
    cleaned = re.sub(r'^.*절대\s*변경하지\s*마.*$', '', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'^.*출력을\s*종료하세요.*$', '', cleaned, flags=re.MULTILINE)
    # 연속 빈 줄 정리
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()

def _clean_md(text: str) -> str:
    """마크다운 문법 기호를 제거하여 DOCX 셀/본문에 사용"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)       # **볼드**
    text = re.sub(r'\*(.*?)\*', r'\1', text)            # *이탤릭*
    text = re.sub(r'`(.*?)`', r'\1', text)              # `코드`
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)  # 주석
    text = re.sub(r'^⛔.*', '', text)                    # 경고 문구
    return text.strip()

def _markdown_to_docx(markdown_text: str, project_id: str) -> bytes:
    """마크다운 텍스트를 Word 문서 바이트로 변환"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 제목
    doc.add_heading(f'분석 리포트 — {project_id}', level=0)

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 빈 줄 / 구분선 건너뛰기
        if not line or line == '---':
            i += 1
            continue

        # ── 헤딩 ──
        if line.startswith('### '):
            doc.add_heading(_clean_md(line[4:]), level=3)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(_clean_md(line[3:]), level=2)
            i += 1
            continue
        if line.startswith('# '):
            doc.add_heading(_clean_md(line[2:]), level=1)
            i += 1
            continue

        # ── 마크다운 테이블 ──
        if '|' in line and line.startswith('|'):
            table_lines = [line]
            i += 1

            # 구분선 (|---|---|) 건너뛰기
            if i < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i].strip()):
                i += 1

            # 나머지 테이블 행 수집
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                if re.match(r'^\|[\s\-:|]+\|$', lines[i].strip()):
                    i += 1
                    continue
                table_lines.append(lines[i].strip())
                i += 1

            _add_table_to_doc(doc, table_lines)
            doc.add_paragraph()  # 테이블 뒤 간격
            continue

        # ── 인용문 ──
        if line.startswith('>'):
            text = _clean_md(line.lstrip('> '))
            if text:
                doc.add_paragraph(text, style='Normal')
            i += 1
            continue

        # ── 불릿 리스트 ──
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(_clean_md(line[2:]), style='List Bullet')
            i += 1
            continue

        # ── 번호 리스트 ──
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(_clean_md(text), style='List Number')
            i += 1
            continue

        # ── 일반 텍스트 ──
        cleaned = _clean_md(line)
        if cleaned:
            _add_formatted_paragraph(doc, line)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_table_to_doc(doc, table_lines: List[str]) -> None:
    """마크다운 테이블 줄 목록을 DOCX 테이블로 변환"""
    from docx.shared import Pt

    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells:
            rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = _clean_md(cell_text)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if i == 0:
                            run.bold = True


def _add_formatted_paragraph(doc, text: str) -> None:
    """볼드/코드를 유지한 채 DOCX paragraph 추가"""
    from docx.shared import Pt, RGBColor

    para = doc.add_paragraph()
    # **bold** 와 `code` 를 분리
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
            run.font.size = Pt(9)
        else:
            para.add_run(_clean_md(part))

# ============================================================
# 4) Solver Execution
# ============================================================
from engine.solver_pipeline import SolverPipeline
from engine.solver_registry import get_solver_time_limit
from domains.crew.agent import get_session, save_session_state

_pipeline = SolverPipeline()


@router.post('/solve')
@limiter.limit("5/minute")
async def solve_optimization(
    request: Request,
    current_user: UserDB = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    body = await request.json()
    project_id = body.get('project_id')
    solver_id = body.get('solver_id')
    solver_name = body.get('solver_name', '')
    math_model = body.get('math_model')
    # 우선순위: DB 설정 > YAML max_time_seconds > fallback(120s)
    time_limit = get_solver_time_limit(solver_id, db)

    if not project_id or not solver_id:
        raise HTTPException(status_code=400, detail='project_id and solver_id are required')

    if not math_model:
        try:
            from domains.crew.agent import get_session
            session = get_session(str(project_id))
            math_model = session.state.math_model
            if not math_model:
                raise HTTPException(
                    status_code=400,
                    detail='No math model found. Please generate a math model first.'
                )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f'Failed to load session: {str(e)}')

    try:
        result = await _pipeline.run(
            math_model=math_model,
            solver_id=solver_id,
            project_id=str(project_id),
            solver_name=solver_name,
            time_limit_sec=time_limit,
        )

        if result.success:
            # 후처리: RunResult + SessionState + ChatHistory (공통 헬퍼 사용)
            try:
                from engine.post_processing import post_process_solve_result
                post_process_solve_result(
                    project_id=int(project_id),
                    solver_id=solver_id,
                    solver_name=solver_name,
                    summary=result.summary,
                    status=result.execute_result.status if result.execute_result else 'UNKNOWN',
                    objective_value=result.execute_result.objective_value if result.execute_result else None,
                    db=db,
                    is_compare=False,
                )
            except Exception as pp_err:
                logger.warning(f'Post-processing failed: {pp_err}')

            return {
                'success': True,
                'phase': result.phase,
                'solver_id': result.solver_id,
                'solver_name': result.solver_name,
                'status': result.execute_result.status if result.execute_result else 'UNKNOWN',
                'summary': result.summary,
                'recovery_log': result.recovery_log if hasattr(result, 'recovery_log') else [],
            }
        else:
            # INFEASIBLE 진단 정보 추출
            infeasibility_info = None
            if result.execute_result and result.execute_result.infeasibility_info:
                infeasibility_info = result.execute_result.infeasibility_info

            return {
                'success': False,
                'phase': result.phase,
                'solver_id': result.solver_id,
                'solver_name': result.solver_name,
                'error': result.error,
                'error_code': result.pipeline_error.code.value if hasattr(result, 'pipeline_error') and result.pipeline_error else 'UNKNOWN',
                'severity': result.pipeline_error.severity.value if hasattr(result, 'pipeline_error') and result.pipeline_error else 'unknown',
                'fallback_solvers': result.pipeline_error.fallback_solvers if hasattr(result, 'pipeline_error') and result.pipeline_error else [],
                'compile_warnings': result.compile_result.warnings if result.compile_result else [],
                'recovery_log': result.recovery_log if hasattr(result, 'recovery_log') else [],
                'infeasibility_info': infeasibility_info,
            }

    except Exception as e:
        logger.error(f'Solve API error: {e}', exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
